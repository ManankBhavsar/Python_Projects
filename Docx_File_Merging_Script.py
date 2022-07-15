import os
import re
import docx
import logging, socket

logging.basicConfig(filename='m1.log', level=logging.DEBUG, format="%(asctime)s" + "-" +
                                                                   socket.gethostname() + "-" +
                                                                   socket.gethostbyname(socket.gethostname()) +
                                                                   "-" +
                                                                   "%(levelname)s %(message)-16s",
                    datefmt="%d-%m-%Y %H:%M")

path = input("Enter path of directory :")  # variable
try:
    count1 = 1
    count2 = 1
    count3 = 1
    os.makedirs(os.path.join('E:\inuron\Python', 'Testing1'), exist_ok=True)
    path1 = 'E:\inuron\Python\Testing1'

    for i, j in enumerate(os.listdir(path)):

        if re.search('\d*\s[(]\d[)][.]docx$', j):
            os.rename(os.path.join(path, j), os.path.join(path1, "Advanced_Assignment_" + str(count1) + '.docx'))
            doc = docx.Document(os.path.join(path1, "Advanced_Assignment_" + str(count1) + '.docx'))
            if count1 == 1:
                doc1 = docx.Document()
            else:
                doc1 = docx.Document('E:\inuron\Python\Advanced_Assignment.docx')
            doc1.add_heading("Advanced_Assignment_" + str(count1), 0)
            for para in doc.paragraphs:
                if re.findall('\S',para.text) != 0:
                    doc1.add_paragraph(para.text)
            doc1.save('E:\inuron\Python\Advanced_Assignment.docx')
            count1 += 1

        elif re.search('_\d*[.]docx$', j):
            os.rename(os.path.join(path, j), os.path.join(path1, "Basic_Assignment_" + str(count2) + '.docx'))

            doc = docx.Document(os.path.join(path1, "Basic_Assignment_" + str(count2) + '.docx'))
            if count2 == 1:
                doc1 = docx.Document()
            else:
                doc1 = docx.Document('E:\inuron\Python\Basic_Assignment.docx')
            doc1.add_heading("Basic_Assignment_" + str(count2), 0)
            for para in doc.paragraphs:
                if re.findall('\S',para.text) != 0:
                    doc1.add_paragraph(para.text)
            doc1.save('E:\inuron\Python\Basic_Assignment.docx')

            count2 += 1

        elif re.search('[A-Za-z]\d*[.]docx$', j):
            os.rename(os.path.join(path, j), os.path.join(path1, "Practical_Assignment_" + str(count3) + '.docx'))

            doc = docx.Document(os.path.join(path1, "Practical_Assignment_" + str(count3) + '.docx'))
            if count3 == 1:
                doc1 = docx.Document()
            else:
                doc1 = docx.Document('E:\inuron\Python\Practical_Assignment.docx')
            doc1.add_heading("Practical_Assignment_" + str(count3), 0)
            for para in doc.paragraphs:
                if re.findall('\S', para.text) != 0:
                    doc1.add_paragraph(para.text)
            doc1.save('E:\inuron\Python\Practical_Assignment.docx')

            count3 += 1

        else:
            raise Exception("No docx files found")
except Exception as e:
    logging.info(e)
